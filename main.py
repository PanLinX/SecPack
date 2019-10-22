#coding=utf-8


import sys
import os
import win32com.client
import time
import logging
from docx import Document
from docx.shared import Cm
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import PIL
from win32com.client import Dispatch, constants, gencache
from pdf2image import convert_from_path
from wand.image import Image
from conf_test import *
import datetime

memo = {}

class FileProcessTool(object):


    def CQFilePro(self):
        """"""
        document = Document(FILE1)
        paragraphs = document.paragraphs
        #paragraphs[3].style.font.name = u"宋体"
        #text = re.sub('2018-7-26','2018-8-20',paragraphs[3].text)
        #paragraphs[3].text = text
        print("开始修改文件")
        paragraphs[3].text = r"编号：QA/PCD.ELF-"+DATE+r"-01"
        tables = document.tables
        tables[0].cell(0,8).text = r"_"+VERSION
        tables[0].cell(2,8).text = DATE_1
        tables[0].cell(4,2).text = MISSON_NO + MISSON_NAME
        tables[0].cell(8,2).text = DESCRIPTION
        tables[0].cell(10,2).text = PICI
        if isProblem == 1 and UAT_NO is not "":
            tables[0].cell(5,2).text = UAT_NO
        #time.sleep(0.5)
        print("修改文件结束")
        document.save(FILE1)

    def InstallManu(self):
        """"""
        document = Document(FILE2)
        #for ele in document.styles:
        #    print("ele: %s" % ele)
        paragraphs = document.paragraphs
        print("开始修改文件2")
        paragraphs[10].text = r"日    期："+DATE_1
        paragraphs[10].paragraph_format.left_indent = Cm(0.7*4)
        paragraphs[11].text = r"文档编号: _"+VERSION
        paragraphs[11].paragraph_format.left_indent = Cm(0.7*4)
        paragraphs[11].style = document.styles['样式1']
        paragraphs[11].style.font.size = Cm(0.58)
        #style = document.styles['Normal']
        #font = style.font
        #font.name = '宋体'
        #font.size = Cm(0.58)
        tables = document.tables
        tables[0].cell(1,0).text = DATE
        #time.sleep(0.5)
        print("修改文件2结束")
        paragraphs[59].text = STRING2
        document.save(FILE2)


    def ImageCut(self,filename,outputfile):
        """"""
        filename = FILE5
        outputfile = FILE6
        img = PIL.Image.open(filename)
        w, h = img.size
        img.crop((45, 80, w-45, h-71)).save(outputfile)

    def word2pdf(self,filename):
        #input = filename + '.docx'
        #output = filename + '.pdf'
        #pdf_name = output
        input = FILE1
        output = FILE3
        pdf_name = FILE3
        try:
            print("开始转换",input,"------->>>>>",output)
            gencache.EnsureModule('{00020905-0000-0000-C000-000000000046}', 0, 8, 4)
            w = Dispatch("Word.Application")
            try:
                doc = w.Documents.Open(input, ReadOnly=1)
                doc.ExportAsFixedFormat(output, constants.wdExportFormatPDF,\
                                        Item=constants.wdExportDocumentWithMarkup,\
                                        CreateBookmarks=constants.wdExportCreateHeadingBookmarks)
            except:
                print("exception1")
            finally:
                w.Quit(constants.wdDoNotSaveChanges)

            if os.path.isfile(pdf_name):
                print("translate success 转换结束。")
                return True
            else:
                print('translate fail')
                return False
        except:
            print('exception2')
            return -1

    def word2pdf111(self,filename):
        #input = filename + '.docx'
        #output = filename + '.pdf'
        #pdf_name = output
        input = FILE1
        output = FILE3
        pdf_name = FILE3
        print("开始转换",input,"------->>>>>",output)
        gencache.EnsureModule('{00020905-0000-0000-C000-000000000046}', 0, 8, 4)
        w = Dispatch("Word.Application")
        doc = w.Documents.Open(input, ReadOnly=1)
        doc.ExportAsFixedFormat(output, constants.wdExportFormatPDF,\
                                        Item=constants.wdExportDocumentWithMarkup,\
                                        CreateBookmarks=constants.wdExportCreateHeadingBookmarks)
        w.Quit(constants.wdDoNotSaveChanges)

        if os.path.isfile(pdf_name):
            print("translate success 转换结束。")
            return True
        else:
            print('translate fail')
            return False

    def NameChange(self):
        """"""
        print("开始修改名字")
        RENAME_FILE1 = BASE_DIR + r"\文档\软件_" + VERSION + r").docx"
        RENAME_BASE_DIR = r"D:\Datas\PyCharm\FileProcess\B_" + VERSION
        os.rename(FILE1, RENAME_FILE1)
        os.rename(BASE_DIR, RENAME_BASE_DIR)
        print("修改名字结束")

    def pdf2jpg(self):
        pdf_file = FILE3
        output = FILE4
        with Image(filename=pdf_file) as img:
            img.save(filename=output)

    def InsertJpgFile(self):
        insertjpg = FILE6
        document = Document(FILE1)
        paragraphs = document.paragraphs
        paragraphs[0].text = ""
        run = paragraphs[0].add_run()
        run.add_picture(insertjpg, width=Inches(6))
        document.save(FILE1)

    def DeleteJpgFile(self):
        document = Document(FILE1)
        tables = document.tables
        tables[0].cell(15, 2).text = ""
        tables[0].cell(15, 8).text = ""
        document.save(FILE1)

if __name__ == '__main__':
    starttime = datetime.datetime.now()
    fp = FileProcessTool()
    #1、替换软件产品出入库申请单必要信息
    fp.CQFilePro()

    #2、替换安装手册必要信息
    fp.InstallManu()

    #3、软件产品出入库申请单word转pdf
    fp.word2pdf('filename')

    #4、软件产品出入库申请单pdf转jpg
    fp.pdf2jpg()

    #5、软件产品出入库申请单第二页jpg切割
    fp.ImageCut('file1','file2')

    #6、软件产品出入库申请单插入jpg图片
    fp.InsertJpgFile()

    #7、软件产品出入库申请单删除多余签名信息
    fp.DeleteJpgFile()

    #8、修改软件产品出入库申请单及文件夹名字
    fp.NameChange()
    endtime = datetime.datetime.now()
    print(endtime - starttime)